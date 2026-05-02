"""
generate-docx.py

Generates a formatted episode script Word document from a JSON episode file,
using Episode_Script.docx as the style template.

Usage:
  python generate-docx.py                        # reads episode.json
  python generate-docx.py my-episode.json        # reads a specific file

See references/episode-template.json for the expected JSON structure.

Styles inherited from Episode_Script.docx:
  Heading 1      - section headings (INTRODUCTION, TOPIC DISCUSSION, etc.)
  Heading 2      - subsection headings with time markers
  Normal         - body paragraphs, each with a bold label and italic content

Paragraph format for body content:
  LABEL  [italic instruction or generated text]
  The label is bold; the content following it is italic.

Spacing conventions (reproduced from Episode_Script.docx):
  After Heading 1            : 2 blank Normal paragraphs
  After Heading 2            : 0 blank Normal paragraphs
  Between content paragraphs : 2 blank Normal paragraphs
  Before any heading         : 1 additional blank (3 total after last content)
  After REFERENCE MATERIALS  : 1 blank Normal paragraph
"""

import json
import sys
from pathlib import Path

from docx import Document

TEMPLATE = Path(__file__).parent.parent / "assets" / "Episode_Script.docx"


# ── JSON LOADER ───────────────────────────────────────────────────────────────

def load_episode(path: Path) -> dict:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


# ── DOCUMENT HELPERS ──────────────────────────────────────────────────────────

def clear_body(doc):
    """Remove all paragraphs/tables from the template body, keeping w:sectPr."""
    body = doc.element.body
    for child in list(body):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)


# Style names in this template are title-case ('Heading 1'), so dict-style
# lookup via python-docx's BabelFish fails. Resolve by iterating instead.
_style_cache: dict = {}

def resolve_style(doc, name):
    if name not in _style_cache:
        for s in doc.styles:
            if s.name == name:
                _style_cache[name] = s
                break
        else:
            raise KeyError(f"Style not found in template: {name!r}")
    return _style_cache[name]


def para(doc, text, style_name):
    p = doc.add_paragraph(text)
    p.style = resolve_style(doc, style_name)
    return p


def blank(doc):
    return para(doc, "", "Normal")


def labeled_para(doc, label, text):
    """Write a Normal paragraph with a bold label and italic content."""
    p = doc.add_paragraph()
    p.style = resolve_style(doc, "Normal")
    p.add_run(f"{label}  ").bold = True
    p.add_run(text).italic = True
    return p


# ── DOCUMENT BUILDER ──────────────────────────────────────────────────────────

def write_content_paragraphs(doc, paragraphs):
    """Write labeled content paragraphs with 2 blank lines between each."""
    for i, p in enumerate(paragraphs):
        labeled_para(doc, p["label"], p["text"])
        blank(doc)
        blank(doc)


def build_document(ep: dict):
    _style_cache.clear()
    doc = Document(TEMPLATE)
    clear_body(doc)

    for s_idx, section in enumerate(ep["sections"]):
        if s_idx > 0:
            blank(doc)  # 3rd blank before heading (2 already written after last content)

        para(doc, section["heading"], "Heading 1")

        if "paragraphs" in section:
            blank(doc)
            blank(doc)  # 2 blanks after Heading 1
            write_content_paragraphs(doc, section["paragraphs"])

        if "subsections" in section:
            blank(doc)
            blank(doc)  # 2 blanks after Heading 1 before first Heading 2

            for sub_idx, sub in enumerate(section["subsections"]):
                if sub_idx > 0:
                    blank(doc)  # 3rd blank before Heading 2

                h2_text = f'{sub["topic_subheading"]}  —  [{sub["time_marker"]}]'
                para(doc, h2_text, "Heading 2")
                # 0 blanks after Heading 2 — content follows immediately
                write_content_paragraphs(doc, sub["paragraphs"])

    # Reference Materials section
    blank(doc)  # 3rd blank before heading
    para(doc, "REFERENCE MATERIALS", "Heading 1")
    blank(doc)  # 1 blank after this heading (matches Episode_Script.docx)
    para(doc, "List all books, commentaries, and resources referenced in this episode:", "Normal")
    blank(doc)
    blank(doc)
    for ref in ep["references"]:
        para(doc, ref, "Normal")
    blank(doc)
    blank(doc)

    p = doc.add_paragraph()
    p.style = resolve_style(doc, "Normal")
    p.add_run("✔  Ready for recording. Episode is approved and production-ready.").bold = True

    out = Path.cwd() / ep["output"]
    doc.save(out)
    print(f"Saved: {out}")


# ── ENTRY POINT ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) > 1:
        json_path = Path(sys.argv[1])
    else:
        json_path = Path(__file__).parent / "episode.json"

    if not json_path.exists():
        print(f"Error: JSON file not found — {json_path}")
        print("Usage: python generate-docx.py [episode.json]")
        sys.exit(1)

    build_document(load_episode(json_path))
