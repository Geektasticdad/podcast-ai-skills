"""
generate-docx.py

Generates a formatted series episode narrative Word document from a JSON episode file,
using Episode TBD.docx as the style template.

Usage:
  python generate-docx.py                          # reads series.json
  python generate-docx.py my-episode.json          # reads a specific file

See references/series-template.json for the expected JSON structure.

Styles inherited from Episode TBD.docx:
  Title          - episode title      (Arial 26pt, black)
  Heading 1      - section headings   (Arial 20pt, black)
  Heading 2      - subsection heads   (Arial 16pt, black, spaced)
  Normal         - body paragraphs    (Arial 11pt, ~14pt after, 1.38x leading)
  List Paragraph - reference items
"""

import json
import sys
from pathlib import Path

from docx import Document

TEMPLATE = Path(__file__).parent.parent / "assets" / "Episode TBD.docx"


# ── JSON LOADER ───────────────────────────────────────────────────────────────

def load_episode(path: Path) -> dict:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


# ── DOCUMENT BUILDER ──────────────────────────────────────────────────────────

def clear_body(doc):
    """Remove all paragraphs/tables from the template body, keeping w:sectPr."""
    body = doc.element.body
    for child in list(body):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)


# This template stores style names in title case ('Heading 1') rather than
# the lowercase form python-docx's BabelFish expects, so dict-style lookup
# (doc.styles['Heading 1']) fails. We resolve styles by iterating instead.
_style_cache: dict = {}

def resolve_style(doc, name):
    if name not in _style_cache:
        for s in doc.styles:
            if s.name == name:
                _style_cache[name] = s
                break
        else:
            raise KeyError(f"Style not found in template: {name!r}")
    return _style_cache[name]


def para(doc, text, style_name):
    p = doc.add_paragraph(text)
    p.style = resolve_style(doc, style_name)
    return p


def build_document(ep: dict):
    _style_cache.clear()
    doc = Document(TEMPLATE)
    clear_body(doc)

    para(doc, ep["title"], "Title")

    for section in ep["sections"]:
        para(doc, section["heading"], "Heading 1")

        if "paragraphs" in section:
            for text in section["paragraphs"]:
                para(doc, text, "Normal")

        if "subsections" in section:
            for sub in section["subsections"]:
                para(doc, sub["heading"], "Heading 2")
                for text in sub["paragraphs"]:
                    para(doc, text, "Normal")

    para(doc, "Reference Materials:", "Heading 1")
    for ref in ep["references"]:
        para(doc, ref, "List Paragraph")

    p = para(doc, "", "Normal")
    p.add_run("Ready for recording. Episode is approved and production-ready.").font.italic = True

    out = Path.cwd() / ep["output"]
    doc.save(out)
    print(f"Saved: {out}")


# ── ENTRY POINT ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) > 1:
        json_path = Path(sys.argv[1])
    else:
        json_path = Path(__file__).parent / "series.json"

    if not json_path.exists():
        print(f"Error: JSON file not found — {json_path}")
        print("Usage: python generate-docx.py [series.json]")
        sys.exit(1)

    build_document(load_episode(json_path))
