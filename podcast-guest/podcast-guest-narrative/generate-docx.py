"""
generate-docx.py

Generates a formatted guest interview Word document from a JSON guest file,
using Episode TBD.docx as the style template.

Usage:
  python generate-docx.py                        # reads guest.json
  python generate-docx.py my-guest.json          # reads a specific file

See guest-template.json for the expected JSON structure.

Sections with type "prose" render as headed paragraphs.
Sections with type "questions" render as numbered questions with italic follow-up prompts.

Styles inherited from Episode TBD.docx:
  Title          - document title     (Arial 26pt, black)
  Heading 1      - section headings   (Arial 20pt, black)
  Normal         - body paragraphs and questions (Arial 11pt, ~14pt after, 1.38x leading)
  List Paragraph - reference items
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

TEMPLATE = Path(__file__).parent / "Episode TBD.docx"


# ── JSON LOADER ───────────────────────────────────────────────────────────────

def load_guest(path: Path) -> dict:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


# ── DOCUMENT BUILDER ──────────────────────────────────────────────────────────

def clear_body(doc):
    """Remove all paragraphs/tables from the template body, keeping w:sectPr."""
    body = doc.element.body
    for child in list(body):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)


# This template stores style names in title case ('Heading 1') rather than
# the lowercase form python-docx's BabelFish expects, so dict-style lookup
# (doc.styles['Heading 1']) fails. We resolve styles by iterating instead.
_style_cache: dict = {}

def resolve_style(doc, name):
    if name not in _style_cache:
        for s in doc.styles:
            if s.name == name:
                _style_cache[name] = s
                break
        else:
            raise KeyError(f"Style not found in template: {name!r}")
    return _style_cache[name]


def para(doc, text, style_name):
    p = doc.add_paragraph(text)
    p.style = resolve_style(doc, style_name)
    return p


def add_questions(doc, questions):
    """Add numbered questions with indented italic follow-up prompts."""
    for i, q in enumerate(questions, 1):
        qp = doc.add_paragraph(f"{i}. {q['question']}")
        qp.style = resolve_style(doc, "Normal")

        fp = doc.add_paragraph()
        fp.style = resolve_style(doc, "Normal")
        fp.paragraph_format.left_indent = Pt(36)
        fp.add_run(f"Follow-up: {q['followup']}").italic = True


def build_document(guest: dict):
    _style_cache.clear()
    doc = Document(TEMPLATE)
    clear_body(doc)

    para(doc, guest["title"], "Title")

    for section in guest["sections"]:
        para(doc, section["heading"], "Heading 1")

        if section["type"] == "prose":
            for text in section["paragraphs"]:
                para(doc, text, "Normal")
            if "word_count" in section:
                wcp = doc.add_paragraph()
                wcp.style = resolve_style(doc, "Normal")
                wcp.add_run(f"[Word count: ~{section['word_count']} words]").italic = True

        elif section["type"] == "questions":
            add_questions(doc, section["questions"])

    p = para(doc, "", "Normal")
    p.add_run("Ready for recording. Guest interview document is approved and production-ready.").italic = True

    out = Path(__file__).parent / guest["output"]
    doc.save(out)
    print(f"Saved: {out}")


# ── ENTRY POINT ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) > 1:
        json_path = Path(sys.argv[1])
    else:
        json_path = Path(__file__).parent / "guest.json"

    if not json_path.exists():
        print(f"Error: JSON file not found — {json_path}")
        print("Usage: python generate-docx.py [guest.json]")
        sys.exit(1)

    build_document(load_guest(json_path))
